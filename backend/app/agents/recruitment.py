"""Recruitment 도메인 에이전트.

v0.9 확장 포인트:
- 3종 플랫폼 채용공고 자동 생성 (당근알바 / 알바천국 / 사람인) —
  `[JOB_POSTINGS]...[/JOB_POSTINGS]` 마커 안에 세 섹션을 받아 3개 artifact + 부모 세트 1개 저장.
- 채용공고 HTML 포스터 생성 (GPT-4o) — `[POSTING_POSTER_REQUEST]` 마커 혹은
  사용자 발화 휴리스틱. 저장 버킷: `recruitment-posters` (HTML 파일) + artifact.content.
- hiring_drive (채용 기간) artifact 에 `due_label="채용 마감"` + `end_date` 주입 →
  스케쥴러 `scanner.find_date_notifications` 가 D-7/3/1/0 리마인드를 자동 발사.
- profiles.business_type 기반 업종별 CHOICES 분기 (`_recruit_templates.detect_category`).
"""

from __future__ import annotations

import logging
import re

from langsmith import traceable

from app.core.llm import chat_completion
from app.core.supabase import get_supabase
from app.agents.orchestrator import (
    CLARIFY_RULE,
    ARTIFACT_RULE,
)
from app.agents._feedback import feedback_context
from app.agents._suggest import suggest_today_for_domain
from app.agents._artifact import (
    save_artifact_from_reply,
    list_sub_hub_titles,
    pick_sub_hub_id,
    today_context,
    record_artifact_for_focus,
)
from app.agents._recruit_templates import (
    CATEGORY_CHOICES,
    PLATFORM_LABELS,
    POSTING_POSTER_REQUEST_RE,
    VALID_PLATFORMS,
    VALID_TYPES,
    build_recruit_context,
    detect_category,
    detect_recruit_intent,
    parse_platform_sections,
    wants_posting_poster,
)

log = logging.getLogger("boss2.recruitment")


# `save_artifact_from_reply` 가 검증하는 일반 타입 허용 목록.
# (job_posting_set / job_posting_poster 는 전용 파이프라인에서 별도 저장.)
_GENERIC_VALID_TYPES: tuple[str, ...] = (
    "job_posting",
    "interview_questions",
    "interview_evaluation",
    "checklist",
    "guide",
    "hiring_drive",
    "onboarding_checklist",
    "onboarding_plan",
    "education_material",
)

_RESUME_PARSE_SYSTEM = (
    "당신은 이력서 파싱 전문가입니다. "
    "주어진 이력서 텍스트에서 정보를 추출해 JSON만 반환하세요. "
    "없는 정보는 null 또는 빈 배열로 설정하세요. 절대 정보를 추측하거나 만들어내지 마세요.\n\n"
    "분류 기준:\n"
    "- experience: 실제 재직·인턴십 (회사에 소속되어 급여를 받은 경력)\n"
    "- projects: 팀/개인 프로젝트, 해커톤, 사이드 프로젝트, 수업 과제 프로젝트 등\n"
    "- training: 부트캠프, 교육 수료, 연수, 강의 이수 등\n\n"
    "반환 형식 (JSON only, 설명 없이):\n"
    "{\n"
    '  "name": "이름 또는 null",\n'
    '  "phone": "연락처 또는 null",\n'
    '  "email": "이메일 또는 null",\n'
    '  "age": 나이(정수) 또는 null,\n'
    '  "address": "주소 또는 null",\n'
    '  "education": [{"school":"","major":"","degree":"","year":""}],\n'
    '  "experience": [{"company":"","role":"","period":"","description":""}],\n'
    '  "projects": [{"name":"","role":"","period":"","tech_stack":"","description":""}],\n'
    '  "training": [{"institution":"","course":"","period":"","description":""}],\n'
    '  "skills": ["기술1"],\n'
    '  "certifications": ["자격증1"],\n'
    '  "desired_position": "희망직종 또는 null",\n'
    '  "desired_salary": "희망급여 또는 null",\n'
    '  "introduction": "자기소개 전문 또는 null",\n'
    '  "raw_text": "이력서 원문 전체"\n'
    "}"
)

_INTERVIEW_FROM_RESUME_SYSTEM = (
    "당신은 소상공인 채용 전문가입니다. "
    "지원자 이력서를 바탕으로 날카롭고 구체적인 면접 질문을 생성합니다.\n"
    "규칙:\n"
    "- 이력서의 구체적 내용(회사명, 기간, 역할)을 직접 인용해 질문\n"
    "- 경력 공백, 짧은 재직기간, 직무 불일치는 파고드는 질문 포함\n"
    "- 직무 적합성 / 성실성 / 상황 대응력 3축으로 골고루 구성\n"
    "- 번호 목록 형식으로만 답변 (설명 없이 질문만)"
)

_JOB_POSTINGS_RE = re.compile(r"\[JOB_POSTINGS\](.*?)\[/JOB_POSTINGS\]", re.DOTALL)


def _format_resume_table(name: str, a: dict) -> str:
    """파싱된 applicant dict → 마크다운 표 형식 문자열."""
    rows: list[str] = []

    def row(label: str, value: str) -> None:
        if value:
            rows.append(f"| {label} | {value} |")

    row("이름", name)
    row("연락처", a.get("phone") or "")
    row("이메일", a.get("email") or "")
    row("나이", str(a["age"]) if a.get("age") else "")
    row("주소", a.get("address") or "")
    row("희망직종", a.get("desired_position") or "")
    row("희망급여", a.get("desired_salary") or "")

    edu_list = a.get("education") or []
    if edu_list:
        edu_str = " / ".join(
            " ".join(filter(None, [e.get("school"), e.get("major"), e.get("degree"), e.get("year")]))
            for e in edu_list
        )
        row("학력", edu_str)

    skills = a.get("skills") or []
    if skills:
        row("기술스택", ", ".join(skills))

    certs = a.get("certifications") or []
    if certs:
        row("자격증", ", ".join(certs))

    intro = (a.get("introduction") or "").strip()
    if intro:
        row("자기소개", intro[:300] + ("…" if len(intro) > 300 else ""))

    header = f"### {name}\n\n| 항목 | 내용 |\n|---|---|"
    table = header + "\n" + "\n".join(rows)

    exp_list = a.get("experience") or []
    if exp_list:
        lines = ["", "**경력**", "", "| 회사 | 직무 | 기간 | 주요 업무 |", "|---|---|---|---|"]
        for e in exp_list:
            desc = (e.get("description") or "").replace("\n", " ")[:120]
            lines.append(f"| {e.get('company','')} | {e.get('role','')} | {e.get('period','')} | {desc} |")
        table += "\n" + "\n".join(lines)

    proj_list = a.get("projects") or []
    if proj_list:
        lines = ["", "**프로젝트**", "", "| 프로젝트명 | 역할 | 기간 | 기술스택 | 내용 |", "|---|---|---|---|---|"]
        for p in proj_list:
            desc = (p.get("description") or "").replace("\n", " ")[:100]
            lines.append(f"| {p.get('name','')} | {p.get('role','')} | {p.get('period','')} | {p.get('tech_stack','')} | {desc} |")
        table += "\n" + "\n".join(lines)

    training_list = a.get("training") or []
    if training_list:
        lines = ["", "**교육수료**", "", "| 기관 | 과정 | 기간 | 내용 |", "|---|---|---|---|"]
        for t in training_list:
            desc = (t.get("description") or "").replace("\n", " ")[:100]
            lines.append(f"| {t.get('institution','')} | {t.get('course','')} | {t.get('period','')} | {desc} |")
        table += "\n" + "\n".join(lines)

    return table


def suggest_today(account_id: str) -> list[dict]:
    return suggest_today_for_domain(account_id, "recruitment")


SYSTEM_PROMPT_BASE = """당신은 채용 전문 AI 에이전트입니다.
소상공인의 채용공고 작성, 면접 질문 생성, 직원 관리 조언을 담당합니다.

[제공 가능 작업]
- **3종 플랫폼 채용공고 동시 작성** (당근알바·알바천국·사람인) — 플랫폼별 톤·포맷 차이 반영
- 채용공고 **HTML 포스터** 생성 (GPT-4o, 1장) — 사용자가 요청할 때만
- 직무별 면접 질문 세트 생성
- 근로계약 체크리스트 / 온보딩 가이드
- 공채/시즌 채용 기간(hiring_drive) 기획 — start_date/end_date + `due_label="채용 마감"` 필수

[필수 필드 매트릭스 — 모두 확정되기 전엔 아티팩트 블록 금지]
- 공통: 직종/포지션, 근무지(매장명 또는 지역), 고용 형태(정규/파트/알바)
- job_posting / job_posting_set: + 급여(시급/월급/연봉 중 하나 숫자), 주 근무시간, 근무 요일
- interview_questions: + 직무 레벨(신입/경력/시니어), 질문 수
- hiring_drive: + start_date, end_date, 채용 인원

[3종 플랫폼 공고 생성 규약]
사용자가 채용공고를 요청하고 필수 필드가 모두 확정되면, **단일 턴 안에** 아래 블록을 정확히 한 번만 출력하세요:

[JOB_POSTINGS]
title: <세트를 대표하는 한 줄 제목>
sub_domain: Job_posting   # 생략 가능
start_date: YYYY-MM-DD    # 선택 (모집 시작)
end_date:   YYYY-MM-DD    # 선택 (모집 마감)
due_label:  모집 마감       # 선택 (기본: 모집 마감)
---
[당근알바]
<당근알바용 공고 본문 — 친근한 톤, 이모지 1~2개, 300~600자, 📍🕒💰🍴✨ 구조>

[알바천국]
<알바천국용 공고 본문 — 표준 구인 양식, 굵은 헤더 + bullet, 시급을 제목 가까이>

[사람인]
<사람인용 공고 본문 — 공식 톤, 회사소개→담당업무→자격요건→우대→복리후생→근무조건→전형절차 순>
[/JOB_POSTINGS]

중요:
- 3섹션 모두 필수. 빈 섹션/생략 금지.
- `[당근알바]` 헤더 정확히 사용 (다른 플랫폼 명 혼용 금지). `당근마켓`은 `[당근마켓]` 또는 `[당근알바]` 중 전자도 인정되지만 헤더 라벨 하나만.
- 시급 표기는 **2026년 최저임금 10,320원 이상 숫자**로. "협의" 금지.
- 한 턴에 [JOB_POSTINGS] 와 다른 [ARTIFACT] 를 동시 출력하지 마세요.

[채용공고 포스터 생성]
사용자가 "이미지/포스터 만들어줘" 류를 요청하면, 본문 끝에 아래 마커를 정확히 한 번:

[POSTING_POSTER_REQUEST]
platform: karrot        # karrot | albamon | saramin (없으면 karrot)
style:    따뜻한 브라운 톤의 카페 분위기, 미니멀 타이포
[/POSTING_POSTER_REQUEST]

이 마커는 시스템이 GPT-4o 로 standalone HTML 포스터를 생성 후 `job_posting_poster` artifact 로 저장합니다.
에이전트가 HTML 코드를 직접 출력하지 마세요. 본문에선 "포스터를 만들고 있어요" 한 줄만.
포스터 마커를 넣는 턴엔 [ARTIFACT] / [JOB_POSTINGS] 를 함께 넣지 마세요.

[온보딩 자료 생성]
신규 입사자 온보딩 요청 시 아래 3종 중 하나로 저장하세요. 사용자 발화에서 타입을 자동 감지하고,
불명확하면 [CHOICES] 로 먼저 물어보세요.

- onboarding_checklist: 입사 전·당일·첫 주 단계별 체크리스트
- onboarding_plan: 수습 기간 일자별/주별 온보딩 플랜
- education_material: 업무 교육 자료 (매뉴얼·가이드·SOP)

타입 감지 힌트:
- "체크리스트/할 일 목록/준비사항" → onboarding_checklist
- "계획/플랜/일정/단계" → onboarding_plan
- "교육/매뉴얼/가이드/SOP/자료" → education_material

[ARTIFACT] 블록 규약 (온보딩):
- sub_domain: Onboarding  ← 반드시 포함
- type: onboarding_checklist | onboarding_plan | education_material 중 하나
- 업종과 직책을 반영한 실용적 내용 (일반적 항목 + 업종 특화 항목 포함)

""" + ARTIFACT_RULE + CLARIFY_RULE + """

[채용공고 Placeholder 절대 금지 — 최우선]
확정되지 않은 식별 정보(매장명·주소·연락처·가게명·회사명 등) 를 **임의 값으로 채우거나 환각하지 마세요**.
- `[매장명]` · `[주소]` · `[전화번호]` · `[회사명]` 같은 대괄호 플레이스홀더 금지.
- "서울 마포구", "경기도 성남시" 같은 **사용자가 말한 적 없는 임의 지역명** 주입 금지.
- "시급 10,320원" 같은 기본값을 사용자가 언급 없이 자동 적용하지 말 것 — 모르면 물어보세요.

누락된 정보가 있으면 [JOB_POSTINGS]/[ARTIFACT] 블록 출력을 **중단**하고, 한 번에 하나씩 [CHOICES] 로 되물으세요.
우선순위: 급여 > 근무지 주소 > 근무요일/시간 > 고용형태 > 매장명/회사 연락처 > 특이사항.

예시 (직종 불명확):
"채용공고를 만들어드릴게요. 어떤 포지션의 공고인가요?
[CHOICES]
바리스타
홀 서빙
주방 보조
기타 (직접 입력)
[/CHOICES]"
"""


# ──────────────────────────────────────────────────────────────────────────
# profile 조회 (업종 추정용)
# ──────────────────────────────────────────────────────────────────────────
def _get_business_type(account_id: str) -> str | None:
    sb = get_supabase()
    rows = (
        sb.table("profiles")
        .select("business_type")
        .eq("id", account_id)
        .limit(1)
        .execute()
        .data
        or []
    )
    if not rows:
        return None
    bt = (rows[0].get("business_type") or "").strip()
    return bt or None


# ──────────────────────────────────────────────────────────────────────────
# [JOB_POSTINGS] 블록 파싱 + 3종 저장
# ──────────────────────────────────────────────────────────────────────────
def _parse_job_postings_block(reply: str) -> dict | None:
    m = _JOB_POSTINGS_RE.search(reply)
    if not m:
        return None
    inner = m.group(1)
    # meta (첫 `---` 기준 split)
    parts = inner.split("---", 1)
    meta_raw = parts[0]
    body = parts[1] if len(parts) > 1 else inner

    meta: dict[str, str] = {}
    for line in meta_raw.strip().splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            meta[k.strip()] = v.strip()

    sections = parse_platform_sections(body)
    return {"meta": meta, "sections": sections}


def _strip_job_postings_block(reply: str) -> str:
    return _JOB_POSTINGS_RE.sub("", reply).strip()


def _valid_date_or_none(s: str) -> str | None:
    from datetime import date
    s = (s or "").strip()
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", s):
        return None
    try:
        date.fromisoformat(s)
        return s
    except ValueError:
        return None


async def _save_posting_set(account_id: str, parsed: dict) -> str | None:
    """부모 `job_posting_set` + 자식 3개 `job_posting` 저장.

    반환: 부모 artifact_id (실패 시 None).
    """
    meta_in = parsed.get("meta") or {}
    sections = parsed.get("sections") or {}
    title = (meta_in.get("title") or "채용공고 세트").strip()

    metadata: dict = {}
    for k in ("start_date", "end_date", "due_date"):
        v = _valid_date_or_none(meta_in.get(k, ""))
        if v:
            metadata[k] = v
    due_label = (meta_in.get("due_label") or "").strip()
    if due_label:
        metadata["due_label"] = due_label[:200]
    elif "end_date" in metadata or "due_date" in metadata:
        metadata["due_label"] = "모집 마감"

    sb = get_supabase()
    try:
        parent_payload = {
            "account_id": account_id,
            "domains":    ["recruitment"],
            "kind":       "artifact",
            "type":       "job_posting_set",
            "title":      title,
            "content":    "\n\n".join(
                f"## {PLATFORM_LABELS[p]}\n{sections.get(p, '').strip()}"
                for p in VALID_PLATFORMS
                if sections.get(p)
            ),
            "status":     "draft",
            "metadata":   metadata or {},
        }
        res = sb.table("artifacts").insert(parent_payload).execute()
        if not res.data:
            return None
        parent_id = res.data[0]["id"]
        record_artifact_for_focus(parent_id)
    except Exception as exc:
        log.exception("posting_set parent insert failed: %s", exc)
        return None

    # 서브허브 contains 엣지 (best-effort)
    sub_name = (meta_in.get("sub_domain") or "Job_posting").strip()
    hub_id = pick_sub_hub_id(
        sb, account_id, "recruitment",
        prefer_keywords=(sub_name, "Job_posting", "posting", "채용"),
    )
    if hub_id:
        try:
            sb.table("artifact_edges").insert({
                "account_id": account_id,
                "parent_id":  hub_id,
                "child_id":   parent_id,
                "relation":   "contains",
            }).execute()
        except Exception:
            pass

    # activity_logs + embedding (best-effort)
    try:
        sb.table("activity_logs").insert({
            "account_id":  account_id,
            "type":        "artifact_created",
            "domain":      "recruitment",
            "title":       title,
            "description": "채용공고 3종 세트 생성 (당근·알바천국·사람인)",
            "metadata":    {"artifact_id": parent_id},
        }).execute()
    except Exception:
        pass
    try:
        from app.rag.embedder import index_artifact
        await index_artifact(account_id, "recruitment", parent_id, f"{title}\n{parent_payload['content']}")
    except Exception:
        pass

    try:
        from app.memory.long_term import log_artifact_to_memory
        await log_artifact_to_memory(
            account_id, "recruitment", "job_posting_set", title,
            content=parent_payload.get("content"),
            metadata=metadata,
        )
    except Exception:
        pass

    return parent_id


async def _maybe_dispatch_posting_set(account_id: str, reply: str) -> str:
    parsed = _parse_job_postings_block(reply)
    if not parsed:
        return reply
    parent_id = await _save_posting_set(account_id, parsed)
    cleaned = _strip_job_postings_block(reply)
    if parent_id:
        notice = (
            "\n\n---\n"
            f"_(채용공고 3종이 캔버스에 저장되었어요 — 세트 artifact: `{parent_id}`)_"
            "\n\n채용공고 포스터 이미지도 만들어 드릴까요?"
        )
        return cleaned + notice
    return cleaned + "\n\n---\n_(공고 저장 중 오류가 발생했어요. 다시 시도해주세요.)_"


# ──────────────────────────────────────────────────────────────────────────
# [POSTING_POSTER_REQUEST] 블록 처리
# ──────────────────────────────────────────────────────────────────────────
def _parse_poster_marker(reply: str) -> dict | None:
    m = POSTING_POSTER_REQUEST_RE.search(reply)
    if not m:
        return None
    parsed: dict[str, str] = {}
    for line in m.group(1).strip().splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            parsed[k.strip().lower()] = v.strip()
    return parsed


def _strip_poster_marker(reply: str) -> str:
    return POSTING_POSTER_REQUEST_RE.sub("", reply).strip()


def _find_recent_posting_set(account_id: str) -> dict | None:
    sb = get_supabase()
    rows = (
        sb.table("artifacts")
        .select("id,title,content,metadata,created_at")
        .eq("account_id", account_id)
        .eq("kind", "artifact")
        .eq("type", "job_posting_set")
        .order("created_at", desc=True)
        .limit(1)
        .execute()
        .data
        or []
    )
    return rows[0] if rows else None


def _list_posting_sets(account_id: str, limit: int = 10) -> list[dict]:
    sb = get_supabase()
    return (
        sb.table("artifacts")
        .select("id,title,created_at")
        .eq("account_id", account_id)
        .eq("kind", "artifact")
        .eq("type", "job_posting_set")
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
        .data
        or []
    )


async def _maybe_dispatch_poster(account_id: str, reply: str) -> str:
    marker = _parse_poster_marker(reply)
    if not marker:
        return reply
    cleaned = _strip_poster_marker(reply)

    posting_set_id = (marker.get("posting_set_id") or "").strip()
    if not posting_set_id:
        target = _find_recent_posting_set(account_id)
        if not target:
            return cleaned + "\n\n---\n_(아직 저장된 채용공고 세트가 없어요. 공고를 먼저 만들어주세요.)_"
        posting_set_id = target["id"]

    platforms_raw = (marker.get("platforms") or "karrot").lower()
    platforms = [p.strip() for p in platforms_raw.split(",") if p.strip() in VALID_PLATFORMS]
    if not platforms:
        platforms = ["karrot"]

    style = marker.get("style") or ""

    from app.core.poster_gen import generate_job_posting_poster
    results: list[str] = []
    errors: list[str] = []

    for platform in platforms:
        try:
            poster_artifact = await generate_job_posting_poster(
                account_id=account_id,
                posting_set_id=posting_set_id,
                platform=platform,
                style_prompt=style,
            )
            url = poster_artifact.get("public_url") or ""
            link = f"[미리보기]({url})" if url else ""
            results.append(
                f"**{PLATFORM_LABELS[platform]}** — `{poster_artifact['artifact_id']}`  {link}"
            )
        except Exception as exc:
            log.exception("poster generation failed for platform=%s", platform)
            errors.append(f"{PLATFORM_LABELS[platform]}: {str(exc)[:120]}")

    notice = "\n\n---\n"
    if results:
        notice += "_(포스터 생성 완료)_\n" + "\n".join(f"- {r}" for r in results)
    if errors:
        notice += "\n_(실패)_\n" + "\n".join(f"- {e}" for e in errors)

    return cleaned + notice


# ──────────────────────────────────────────────────────────────────────────
# 메인 run
# ──────────────────────────────────────────────────────────────────────────
# ──────────────────────────────────────────────────────────────────────────
# Capability 인터페이스 (function-calling 라우팅용)
# ──────────────────────────────────────────────────────────────────────────
def _fmt_days(days: list[str] | None) -> str:
    return "·".join(days) if days else ""


@traceable(name="recruitment.run_posting_set", run_type="chain")
async def run_posting_set(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    position: str,
    wage_hourly: int | None = None,
    wage_monthly: int | None = None,
    annual_salary: int | None = None,
    weekly_hours: float | None = None,
    work_days: list[str] | None = None,
    work_start: str | None = None,
    work_end: str | None = None,
    employment_type: str | None = None,
    location: str | None = None,
    headcount: int | None = None,
    start_date: str | None = None,
    end_date: str | None = None,
    extra_note: str | None = None,
    business_name: str | None = None,
) -> str:
    """3종 플랫폼 채용공고 동시 작성.

    미확정 식별 정보(매장명·위치·연락처 등) 는 **환각 금지**. 부족하면
    LLM 이 `[CHOICES]` 로 되물어 오케스트레이터 단까지 돌려보내게 한다.
    """
    lines: list[str] = [f"[직종] {position}"]
    if business_name:
        lines.append(f"[매장명] {business_name}")
    if employment_type:
        lines.append(f"[고용형태] {employment_type}")
    if headcount:
        lines.append(f"[모집인원] {headcount}명")
    if location:
        lines.append(f"[근무지] {location}")
    if wage_hourly:
        lines.append(f"[시급] {wage_hourly:,}원")
    if wage_monthly:
        lines.append(f"[월급] {wage_monthly:,}원")
    if annual_salary:
        lines.append(f"[연봉] {annual_salary:,}원")
    if weekly_hours:
        lines.append(f"[주 근무시간] {weekly_hours}시간")
    if work_days:
        lines.append(f"[근무요일] {_fmt_days(work_days)}")
    if work_start and work_end:
        lines.append(f"[근무시간대] {work_start}~{work_end}")
    if start_date:
        lines.append(f"[모집 시작] {start_date}")
    if end_date:
        lines.append(f"[모집 마감] {end_date}")
    if extra_note:
        lines.append(f"[특이사항] {extra_note}")

    # 핵심 정보 누락 목록 — placeholder 환각 방지용으로 LLM 에 명시 전달
    missing: list[str] = []
    if not business_name:
        missing.append("매장명/상호")
    if not (wage_hourly or wage_monthly or annual_salary):
        missing.append("급여(시급/월급/연봉)")
    if not location:
        missing.append("근무지(매장 주소)")
    if not work_days:
        missing.append("근무 요일")
    if not work_start or not work_end:
        missing.append("근무 시작·종료 시각")
    if not employment_type:
        missing.append("고용 형태(정규·계약·알바)")
    if not headcount:
        missing.append("모집 인원")
    if not start_date:
        missing.append("모집 시작일")
    if not end_date:
        missing.append("모집 마감일")

    if missing:
        synthetic = (
            "사용자가 채용 공고 작성을 요청했는데 아래 필수 정보가 아직 확정되지 않았습니다:\n"
            f"- 누락: {', '.join(missing)}\n\n"
            "확정된 정보:\n"
            + "\n".join(lines)
            + "\n\n**반드시 [CHOICES] 로 누락 항목 중 가장 근본적인 것 하나를 되물으세요.** "
            "[JOB_POSTINGS] 또는 [ARTIFACT] 블록을 절대 출력하지 말고, "
            "`[매장명]` · `[주소]` · `[전화번호]` 같은 placeholder 로 채운 임의 공고도 만들지 마세요. "
            "한 번에 하나만 물어보세요.\n\n"
            f"원본 사용자 요청: {message}"
        )
    else:
        synthetic = (
            "채용공고를 당근알바·알바천국·사람인 3종 플랫폼으로 만들어주세요. "
            "아래 조건이 모두 확정되었으니 추가 질문 없이 바로 [JOB_POSTINGS] 블록을 출력하세요.\n"
            + "\n".join(lines)
            + f"\n\n원본 사용자 요청: {message}"
        )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


@traceable(name="recruitment.run_posting_poster", run_type="chain")
async def run_posting_poster(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    posting_set_id: str | None = None,
    platforms: list[str] | None = None,
    style: str = "",
) -> str:
    """선택한 job_posting_set 을 근거로 HTML 포스터 생성 (복수 플랫폼 지원)."""
    valid_platforms = [p for p in (platforms or []) if p in VALID_PLATFORMS]

    # 공고 선택이 필요한 경우
    if not posting_set_id:
        posting_sets = _list_posting_sets(account_id)
        if not posting_sets:
            return "아직 저장된 채용공고가 없어요. 공고를 먼저 작성해주세요."
        if len(posting_sets) == 1:
            posting_set_id = posting_sets[0]["id"]
        else:
            items = "\n".join(
                f"- [{ps['title']}] (작성일: {ps['created_at'][:10]}, id: {ps['id']})"
                for ps in posting_sets
            )
            synthetic = (
                f"사용자가 채용공고 포스터 생성을 요청했습니다. 저장된 공고 목록:\n{items}\n\n"
                "[CHOICES] 로 어느 공고로 포스터를 만들지 선택하게 해주세요. "
                "목록에서 번호나 제목으로 선택할 수 있다고 안내하세요.\n\n"
                f"원본 요청: {message}"
            )
            return await run(synthetic, account_id, history, rag_context, long_term_context)

    # 플랫폼 선택이 필요한 경우
    if not valid_platforms:
        synthetic = (
            "사용자가 채용공고 포스터 생성을 요청했습니다. 플랫폼을 선택해야 합니다.\n"
            "[CHOICES] 로 당근알바·알바천국·사람인 중 선택하게 해주세요. "
            "복수 선택도 가능하다고 안내하세요 (예: '당근알바 + 알바천국').\n\n"
            f"원본 요청: {message}"
        )
        return await run(synthetic, account_id, history, rag_context, long_term_context)

    # 포스터 생성 마커 출력
    platforms_str = ",".join(valid_platforms)
    synthetic = (
        "채용공고 포스터를 생성해주세요. 아래 마커를 그대로 출력하세요:\n\n"
        "[POSTING_POSTER_REQUEST]\n"
        f"posting_set_id: {posting_set_id}\n"
        f"platforms: {platforms_str}\n"
        f"style: {style or '깔끔하고 모던한 톤, 따뜻한 브라운 계열'}\n"
        "[/POSTING_POSTER_REQUEST]"
    )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


@traceable(name="recruitment.run_interview", run_type="chain")
async def run_interview(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    position: str,
    level: str = "신입",
    count: int = 5,
) -> str:
    """직종·레벨별 면접 질문 세트 생성."""
    synthetic = (
        f"{position} 직무의 {level} 지원자용 면접 질문 {count}개를 생성해주세요. "
        "추가 질문 없이 바로 [ARTIFACT] 블록을 type=interview_questions, sub_domain=Interviews 로 저장하세요.\n\n"
        f"원본 사용자 요청: {message}"
    )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


@traceable(name="recruitment.run_resume_parse", run_type="chain")
async def run_resume_parse(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
) -> str:
    """구직자 이력서 파일(복수 가능)을 파싱해 resumes 테이블에 저장."""
    import json as _json
    from app.agents._upload_context import get_pending_upload, get_pending_uploads

    uploads = get_pending_uploads() or []
    if not uploads:
        single = get_pending_upload()
        if single:
            uploads = [single]

    if not uploads:
        return (
            "이력서 파일을 첨부해주세요. 파일을 채팅창에 업로드한 후 다시 요청해주세요.\n\n"
            "[CHOICES]\n이력서 파일 업로드할게요\n[/CHOICES]"
        )

    sb = get_supabase()
    saved: list[dict] = []

    for up in uploads:
        content = (up.get("content") or "").strip()
        file_name = up.get("original_name") or up.get("title") or "이력서"
        if not content:
            log.warning("[resume_parse] empty content for file=%s", file_name)
            continue

        parse_resp = await chat_completion(
            messages=[
                {"role": "system", "content": _RESUME_PARSE_SYSTEM},
                {"role": "user", "content": f"다음 이력서를 파싱해주세요:\n\n{content[:6000]}"},
            ],
            model="gpt-4o",
            response_format={"type": "json_object"},
        )
        raw_json = parse_resp.choices[0].message.content or "{}"
        try:
            applicant = _json.loads(raw_json)
        except Exception:
            applicant = {"raw_text": content}

        applicant["raw_text"] = applicant.get("raw_text") or content

        try:
            row = (
                sb.table("resumes")
                .insert({
                    "account_id": account_id,
                    "file_name": file_name,
                    "applicant": applicant,
                })
                .execute()
                .data
            )
        except Exception:
            log.warning("[resume_parse] DB insert failed for file=%s", file_name)
            continue
        if row:
            saved.append({
                "id": row[0]["id"],
                "name": (applicant.get("name") or "").strip() or file_name,
                "applicant": applicant,
            })
        else:
            log.warning("[resume_parse] DB insert returned no data for file=%s", file_name)

    if not saved:
        return "이력서 파싱에 실패했습니다. 파일이 텍스트를 포함하는지 확인해주세요."

    # 파싱 완료 → 장기기억에 지원자 정보 누적
    try:
        from app.memory.long_term import log_artifact_to_memory
        for s in saved:
            a = s["applicant"]
            mem_lines = [f"지원자 {s['name']} 이력서 파싱 완료."]
            if a.get("desired_position"):
                mem_lines.append(f"희망직종: {a['desired_position']}.")
            exp = a.get("experience") or []
            if exp:
                mem_lines.append(f"경력: {', '.join(e.get('company','') for e in exp[:3])}.")
            proj = a.get("projects") or []
            if proj:
                mem_lines.append(f"프로젝트: {', '.join(p.get('name','') for p in proj[:3])}.")
            skills = a.get("skills") or []
            if skills:
                mem_lines.append(f"주요 기술: {', '.join(skills[:6])}.")
            await log_artifact_to_memory(
                account_id, "recruitment", "resume_parse", f"{s['name']} 이력서",
                content=" ".join(mem_lines),
                metadata={"resume_id": s["id"]},
            )
    except Exception:
        pass

    summaries: list[str] = []
    for s in saved:
        a = s["applicant"]
        summaries.append(_format_resume_table(s["name"], a))

    summary = "\n\n---\n\n".join(summaries)

    # 사용자가 면접 질문을 원하면 파싱 직후 바로 생성 (2단계 → 1단계 통합)
    interview_kw = ("면접", "질문", "인터뷰", "interview")
    wants_interview = any(kw in message for kw in interview_kw)
    if wants_interview:
        parts = [f"이력서 {len(saved)}건 파싱 완료:\n\n{summary}\n\n---\n"]
        for s in saved:
            questions = await run_resume_interview(
                account_id=account_id,
                message=message,
                history=history,
                long_term_context=long_term_context,
                rag_context=rag_context,
                applicant_name=s["name"],
            )
            parts.append(questions)
        return "\n\n".join(parts)

    choices_items = "\n".join(f"{s['name']} 면접 질문 생성" for s in saved)
    return (
        f"이력서 {len(saved)}건 파싱 완료:\n\n{summary}\n\n"
        f"[CHOICES]\n{choices_items}\n다른 이력서도 올릴게요\n[/CHOICES]"
    )


@traceable(name="recruitment.run_resume_interview", run_type="chain")
async def run_resume_interview(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    applicant_name: str,
    count: int = 7,
) -> str:
    """저장된 이력서를 기반으로 맞춤 면접 질문 생성 후 artifact 저장."""
    sb = get_supabase()
    applicant_name = (applicant_name or "").strip()
    if not applicant_name:
        return "지원자 이름이 필요합니다."

    # account_id 필터 필수 — 최신 파싱 순으로 이름 매칭
    rows = (
        sb.table("resumes")
        .select("*")
        .eq("account_id", account_id)
        .order("parsed_at", desc=True)
        .limit(50)
        .execute()
        .data
        or []
    )
    resume = next(
        (r for r in rows if (r.get("applicant") or {}).get("name") == applicant_name),
        None,
    )
    if resume is None:
        # 이름 정확 매칭 실패 시 파일명으로 폴백
        resume = next(
            (r for r in rows if applicant_name in (r.get("file_name") or "")),
            None,
        )
    if resume is None:
        return f"'{applicant_name}' 이력서를 찾을 수 없습니다. 먼저 이력서를 업로드해주세요."

    applicant = resume.get("applicant") or {}
    resume_id = resume["id"]
    name = (applicant.get("name") or "").strip() or applicant_name

    # 동일 resume_id 에 대해 2분 이내 중복 생성 방지 (planner 이중 dispatch 대응)
    from datetime import datetime, timedelta, timezone as _tz
    cutoff = (datetime.now(_tz.utc) - timedelta(minutes=2)).isoformat()
    try:
        dup = (
            sb.table("artifacts")
            .select("id,content")
            .eq("account_id", account_id)
            .eq("type", "interview_questions")
            .filter("metadata->>resume_id", "eq", resume_id)
            .gte("created_at", cutoff)
            .limit(1)
            .execute()
            .data
        )
        if dup:
            return f"**{name}** 이력서 기반 면접 질문 {count}개 생성 완료.\n\n{dup[0]['content']}"
    except Exception:
        pass

    context_lines = [f"지원자 이름: {name}"]
    if applicant.get("experience"):
        for e in applicant["experience"]:
            context_lines.append(
                f"경력: {e.get('company','')} / {e.get('role','')} / {e.get('period','')} — {e.get('description','')}"
            )
    if applicant.get("education"):
        for ed in applicant["education"]:
            context_lines.append(f"학력: {ed.get('school','')} {ed.get('major','')} {ed.get('year','')}")
    if applicant.get("skills"):
        context_lines.append(f"기술: {', '.join(applicant['skills'])}")
    if applicant.get("certifications"):
        context_lines.append(f"자격증: {', '.join(applicant['certifications'])}")
    if applicant.get("introduction"):
        context_lines.append(f"자기소개: {applicant['introduction'][:500]}")
    if applicant.get("desired_position"):
        context_lines.append(f"희망직종: {applicant['desired_position']}")

    context_text = "\n".join(context_lines)

    resp = await chat_completion(
        messages=[
            {"role": "system", "content": _INTERVIEW_FROM_RESUME_SYSTEM},
            {
                "role": "user",
                "content": (
                    f"아래 지원자 이력서를 바탕으로 날카로운 면접 질문 {count}개를 생성해주세요.\n\n"
                    f"{context_text}"
                ),
            },
        ],
        model="gpt-4o",
    )
    questions_text = (resp.choices[0].message.content or "").strip()

    title = f"{name} 면접 질문"
    payload: dict = {
        "account_id": account_id,
        "domains": ["recruitment"],
        "kind": "artifact",
        "type": "interview_questions",
        "title": title,
        "content": questions_text,
        "status": "draft",
        "metadata": {"resume_id": resume_id},
    }
    try:
        result = sb.table("artifacts").insert(payload).execute()
        if result.data:
            artifact_id = result.data[0]["id"]
            record_artifact_for_focus(artifact_id)
            hub_id = pick_sub_hub_id(sb, account_id, "recruitment", prefer_keywords=("Interviews",))
            if hub_id:
                try:
                    sb.table("artifact_edges").insert({
                        "account_id": account_id,
                        "parent_id": hub_id,
                        "child_id": artifact_id,
                        "relation": "contains",
                    }).execute()
                except Exception:
                    pass
            try:
                sb.table("activity_logs").insert({
                    "account_id": account_id,
                    "type": "artifact_created",
                    "domain": "recruitment",
                    "title": title,
                    "description": "interview_questions 생성됨",
                    "metadata": {"artifact_id": artifact_id, "resume_id": resume_id},
                }).execute()
            except Exception:
                pass
            try:
                from app.memory.long_term import log_artifact_to_memory
                await log_artifact_to_memory(
                    account_id, "recruitment", "interview_questions", title,
                    content=questions_text,
                    metadata={"resume_id": resume_id, "applicant_name": name},
                )
            except Exception:
                pass
    except Exception:
        log.warning("[resume_interview] artifact insert failed for applicant=%s", name)

    return (
        f"**{name}** 이력서 기반 면접 질문 {count}개 생성 완료.\n\n{questions_text}"
    )


@traceable(name="recruitment.run_hiring_drive", run_type="chain")
async def run_hiring_drive(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    title: str,
    start_date: str,
    end_date: str,
    headcount: int,
    target_position: str | None = None,
) -> str:
    """공채/시즌 채용 기간 기획 (스케쥴러 D-리마인드 자동)."""
    pos = target_position or "전 직종"
    synthetic = (
        f"'{title}' 라는 이름의 채용 기간(hiring_drive) artifact 를 만들어주세요. "
        f"기간: {start_date} ~ {end_date}, 모집 인원: {headcount}명, 대상: {pos}.\n"
        "반드시 metadata 에 start_date/end_date/due_label='채용 마감' 을 포함한 [ARTIFACT] 블록으로 저장.\n\n"
        f"원본 사용자 요청: {message}"
    )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


@traceable(name="recruitment.run_checklist_guide", run_type="chain")
async def run_checklist_guide(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    topic: str,
    kind: str = "checklist",
) -> str:
    """채용 관련 체크리스트 또는 가이드."""
    artifact_type = "checklist" if kind == "checklist" else "guide"
    synthetic = (
        f"'{topic}' 주제로 채용 {kind} 를 작성해주세요. "
        f"[ARTIFACT] 블록의 type={artifact_type} 로 저장.\n\n"
        f"원본 사용자 요청: {message}"
    )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


_ONBOARDING_TYPE_LABELS = {
    "onboarding_checklist": "온보딩 체크리스트",
    "onboarding_plan":      "온보딩 플랜",
    "education_material":   "교육 자료",
}


@traceable(name="recruitment.run_onboarding", run_type="chain")
async def run_onboarding(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    position: str,
    onboarding_type: str | None = None,
) -> str:
    """신규 입사자 온보딩 자료 생성 (체크리스트 / 플랜 / 교육자료)."""
    business_type = _get_business_type(account_id) or "소상공인"

    if onboarding_type and onboarding_type in _ONBOARDING_TYPE_LABELS:
        label = _ONBOARDING_TYPE_LABELS[onboarding_type]
        synthetic = (
            f"업종 '{business_type}', 직책 '{position}' 신규 입사자를 위한 {label}을 작성해주세요.\n"
            f"[ARTIFACT] 블록의 type={onboarding_type}, sub_domain=Onboarding 으로 저장.\n"
            "업종 특화 항목을 포함한 실용적 내용으로 작성하세요.\n\n"
            f"원본 사용자 요청: {message}"
        )
    else:
        # 타입 불명확 → LLM이 감지하거나 [CHOICES] 출력
        synthetic = (
            f"업종 '{business_type}', 직책 '{position}' 신규 입사자를 위한 온보딩 자료를 요청합니다.\n"
            "사용자 발화에서 원하는 자료 타입(체크리스트/플랜/교육자료)을 감지하세요.\n"
            "불명확하면 [CHOICES] 로 먼저 물어보고, 명확하면 바로 [ARTIFACT] 블록(sub_domain=Onboarding)으로 저장하세요.\n\n"
            f"원본 사용자 요청: {message}"
        )
    return await run(synthetic, account_id, history, rag_context, long_term_context)


_EVAL_CRITERIA_SYSTEM = """당신은 채용 전문가입니다. 지원자 이력서를 분석해 면접 평가 기준을 JSON으로만 반환하세요.

반환 형식 (JSON only, 설명 없이):
{
  "<평가항목명>": {
    "5": "매우 우수 기준 (구체적으로)",
    "4": "우수 기준",
    "3": "보통 기준",
    "2": "미흡 기준",
    "1": "부적합 기준",
    "checkpoints": ["이력서 기반 확인 포인트1", "이력서 기반 확인 포인트2", "이력서 기반 확인 포인트3"]
  }
}

각 항목의 checkpoints 는 이력서에서 실제로 발견한 구체적 내용을 기반으로 작성하세요."""


def _build_evaluation_docx(
    name: str,
    pos_label: str,
    eval_weights: dict[str, int],
    criteria: dict[str, dict],
) -> bytes:
    """면접 평가표 DOCX 생성. python-docx 사용."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # 페이지 여백 설정
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    def _set_para_format(para, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
        para.alignment = alignment
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)

    def _add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        _set_para_format(p, space_before=12, space_after=6)
        return p

    def _shade_cell(cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # 제목
    title = doc.add_heading("면접 평가표", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph(f"{name}  /  {pos_label}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    for run in sub.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 기본 정보
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = "Table Grid"
    info_cells = info_table.rows[0].cells
    labels = ["면접일", "면접관", "면접 장소", "면접 시간"]
    for i, label in enumerate(labels):
        p = info_cells[i].paragraphs[0]
        p.add_run(f"{label}: ").bold = True
        p.add_run("_" * 10)
        info_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

    # 종합 점수표
    _add_heading("종합 점수표", level=2)
    score_table = doc.add_table(rows=len(eval_weights) + 2, cols=5)
    score_table.style = "Table Grid"

    headers = ["평가 항목", "배점", "점수 (1–5점)", "가중 점수", "메모"]
    header_row = score_table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade_cell(cell, "D9E1F2")

    for row_idx, (cat, pct) in enumerate(eval_weights.items(), start=1):
        row = score_table.rows[row_idx]
        row.cells[0].text = cat
        row.cells[1].text = f"{pct}%"
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ci in [2, 3, 4]:
            row.cells[ci].text = ""

    total_row = score_table.rows[-1]
    total_row.cells[0].text = "합계"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[1].text = "100%"
    total_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade_cell(total_row.cells[0], "FFF2CC")
    _shade_cell(total_row.cells[1], "FFF2CC")
    total_row.cells[3].text = "    / 100점"
    doc.add_paragraph()

    # 역량별 평가 기준
    _add_heading("역량별 평가 기준 및 면접관 코멘트", level=2)

    for cat, pct in eval_weights.items():
        _add_heading(f"{cat}  ({pct}%)", level=3)
        cat_criteria = criteria.get(cat, {})

        if cat_criteria:
            crit_table = doc.add_table(rows=6, cols=2)
            crit_table.style = "Table Grid"
            crit_header = crit_table.rows[0].cells
            crit_header[0].text = "점수"
            crit_header[1].text = "평가 기준"
            for c in crit_header:
                c.paragraphs[0].runs[0].bold = True
                _shade_cell(c, "E2EFDA")

            for i, score in enumerate(["5", "4", "3", "2", "1"], start=1):
                row = crit_table.rows[i]
                row.cells[0].text = f"{score}점"
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[1].text = cat_criteria.get(score, "")

            checkpoints = cat_criteria.get("checkpoints") or []
            if checkpoints:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("이력서 기반 체크포인트").bold = True
                for cp in checkpoints:
                    doc.add_paragraph(f"• {cp}", style="List Bullet")

        comment_p = doc.add_paragraph()
        comment_p.add_run("면접관 코멘트: ").bold = True
        comment_p.add_run("_" * 60)
        doc.add_paragraph()

    # 종합 의견
    _add_heading("종합 의견", level=2)
    for label in ["강점", "우려사항", "종합 평가"]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("_" * 55)
        doc.add_paragraph()

    # 채용 추천도
    rec_p = doc.add_paragraph()
    rec_p.add_run("채용 추천도:  ").bold = True
    rec_p.add_run("□ 강력 추천   □ 추천   □ 보류   □ 비추천")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@traceable(name="recruitment.run_interview_evaluation", run_type="chain")
async def run_interview_evaluation(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    applicant_name: str,
    position: str | None = None,
    custom_categories: list[str] | None = None,
    weights: dict[str, int] | None = None,
) -> str:
    """이력서 기반 맞춤 면접 평가표 생성 → 캔버스 마크다운 저장 (배점 커스터마이징 지원)."""
    import json as _json

    sb = get_supabase()
    rows = (
        sb.table("resumes")
        .select("*")
        .eq("account_id", account_id)
        .order("parsed_at", desc=True)
        .limit(50)
        .execute()
        .data
        or []
    )
    resume = next(
        (r for r in rows if (r.get("applicant") or {}).get("name") == applicant_name),
        None,
    )
    if resume is None:
        resume = next(
            (r for r in rows if applicant_name in (r.get("file_name") or "")),
            None,
        )
    if resume is None:
        return (
            f"'{applicant_name}' 이력서를 찾을 수 없습니다. "
            "먼저 이력서를 업로드해주세요.\n\n"
            "[CHOICES]\n이력서 파일 업로드할게요\n[/CHOICES]"
        )

    applicant = resume.get("applicant") or {}
    resume_id = resume["id"]
    name = (applicant.get("name") or "").strip() or applicant_name
    pos_label = position or applicant.get("desired_position") or "해당 직종"

    # 이력서 컨텍스트 구성
    context_lines = [f"지원자: {name}", f"지원 직종: {pos_label}"]
    for e in applicant.get("experience") or []:
        context_lines.append(
            f"경력: {e.get('company','')} / {e.get('role','')} / {e.get('period','')} — {e.get('description','')}"
        )
    for ed in applicant.get("education") or []:
        context_lines.append(f"학력: {ed.get('school','')} {ed.get('major','')} {ed.get('year','')}")
    if applicant.get("skills"):
        context_lines.append(f"기술: {', '.join(applicant['skills'])}")
    if applicant.get("certifications"):
        context_lines.append(f"자격증: {', '.join(applicant['certifications'])}")
    if applicant.get("introduction"):
        context_lines.append(f"자기소개: {applicant['introduction'][:500]}")
    context_text = "\n".join(context_lines)

    # 배점 결정
    default_categories = {"기술 역량": 50, "태도·성실성": 30, "소통 능력": 20}
    if weights:
        eval_weights = weights
    elif custom_categories:
        per = 100 // len(custom_categories)
        remainder = 100 - per * len(custom_categories)
        eval_weights = {c: per for c in custom_categories}
        first = next(iter(eval_weights))
        eval_weights[first] += remainder
    else:
        eval_weights = default_categories

    categories_desc = ", ".join(f"{c}({p}%)" for c, p in eval_weights.items())

    # LLM으로 평가 기준 JSON 생성
    criteria: dict = {}
    try:
        crit_resp = await chat_completion(
            messages=[
                {"role": "system", "content": _EVAL_CRITERIA_SYSTEM},
                {
                    "role": "user",
                    "content": (
                        f"아래 지원자의 면접 평가 기준을 작성해주세요.\n\n"
                        f"{context_text}\n\n"
                        f"평가 항목: {categories_desc}"
                    ),
                },
            ],
            model="gpt-4o",
            response_format={"type": "json_object"},
        )
        criteria = _json.loads(crit_resp.choices[0].message.content or "{}")
    except Exception:
        log.warning("[interview_evaluation] criteria generation failed for applicant=%s", name)

    # 마크다운 본문 구성 (캔버스 표시용)
    weights_md = "\n".join(
        f"| {cat} | {pct}% | | | |" for cat, pct in eval_weights.items()
    )
    criteria_md_parts = []
    for cat, pct in eval_weights.items():
        cat_c = criteria.get(cat, {})
        cps = "\n".join(f"- {cp}" for cp in (cat_c.get("checkpoints") or []))
        criteria_md_parts.append(
            f"### {cat} ({pct}%)\n"
            + "\n".join(f"- {s}점: {cat_c.get(str(s), '')}" for s in [5, 4, 3, 2, 1])
            + (f"\n\n**체크포인트**\n{cps}" if cps else "")
            + "\n\n**면접관 코멘트:** ___________________________________________"
        )
    criteria_md = "\n\n".join(criteria_md_parts)

    content_md = (
        f"## 면접 평가표 — {name} / {pos_label}\n\n"
        "**면접일:** ______  **면접관:** ______\n\n"
        "### 종합 점수표\n\n"
        "| 평가 항목 | 배점 | 점수(1-5점) | 가중 점수 | 메모 |\n"
        "|---|---|---|---|---|\n"
        f"{weights_md}\n"
        "| **합계** | **100%** | | ___ / 100점 | |\n\n"
        "---\n\n"
        f"{criteria_md}\n\n"
        "---\n\n"
        "### 종합 의견\n\n"
        "**강점:** ___________________________________________\n\n"
        "**우려사항:** ___________________________________________\n\n"
        "**채용 추천도:** □ 강력 추천  □ 추천  □ 보류  □ 비추천"
    )

    title = f"{name} 면접 평가표"
    artifact_meta: dict = {
        "resume_id": resume_id,
        "eval_weights": eval_weights,
        "eval_criteria": criteria,
        "applicant_name": name,
        "position": pos_label,
    }

    try:
        result = sb.table("artifacts").insert({
            "account_id": account_id,
            "domains": ["recruitment"],
            "kind": "artifact",
            "type": "interview_evaluation",
            "title": title,
            "content": content_md,
            "status": "draft",
            "metadata": artifact_meta,
        }).execute()
        if result.data:
            artifact_id = result.data[0]["id"]
            record_artifact_for_focus(artifact_id)
            hub_id = pick_sub_hub_id(sb, account_id, "recruitment", prefer_keywords=("Interviews",))
            if hub_id:
                try:
                    sb.table("artifact_edges").insert({
                        "account_id": account_id,
                        "parent_id": hub_id,
                        "child_id": artifact_id,
                        "relation": "contains",
                    }).execute()
                except Exception:
                    pass
    except Exception:
        log.warning("[interview_evaluation] artifact insert failed for applicant=%s", name)

    return (
        f"**{name}** 면접 평가표를 생성했습니다 ({categories_desc}).\n\n"
        f"{content_md}\n\n"
        "---\n\n"
        "내용을 검토하고 수정하신 후, **DOCX로 저장해줘**라고 말씀하시면 파일로 내보내드릴게요."
    )


@traceable(name="recruitment.run_evaluation_export_docx", run_type="chain")
async def run_evaluation_export_docx(
    *,
    account_id: str,
    message: str,
    history: list[dict],
    long_term_context: str = "",
    rag_context: str = "",
    artifact_id: str | None = None,
    applicant_name: str | None = None,
) -> str:
    """캔버스의 면접 평가표를 DOCX 파일로 변환 후 Supabase Storage에 저장, 다운로드 URL 반환."""
    import uuid as _uuid

    sb = get_supabase()

    # 아티팩트 조회: artifact_id 우선, 없으면 최근 interview_evaluation
    artifact: dict | None = None
    if artifact_id:
        row = (
            sb.table("artifacts")
            .select("*")
            .eq("id", artifact_id)
            .eq("account_id", account_id)
            .maybe_single()
            .execute()
        )
        artifact = row.data
    if not artifact:
        query = (
            sb.table("artifacts")
            .select("*")
            .eq("account_id", account_id)
            .eq("type", "interview_evaluation")
            .order("created_at", desc=True)
            .limit(1)
        )
        if applicant_name:
            query = query.ilike("title", f"%{applicant_name}%")
        rows = query.execute().data or []
        artifact = rows[0] if rows else None

    if not artifact:
        return (
            "저장할 면접 평가표를 찾을 수 없습니다. "
            "먼저 면접 평가표를 생성해주세요."
        )

    meta = artifact.get("metadata") or {}
    name = meta.get("applicant_name") or artifact.get("title", "지원자").replace(" 면접 평가표", "")
    pos_label = meta.get("position") or "해당 직종"
    eval_weights: dict[str, int] = meta.get("eval_weights") or {"기술 역량": 50, "태도·성실성": 30, "소통 능력": 20}
    criteria: dict = meta.get("eval_criteria") or {}

    try:
        docx_bytes = _build_evaluation_docx(name, pos_label, eval_weights, criteria)
    except Exception:
        log.exception("[evaluation_export] docx build failed")
        return "DOCX 파일 생성 중 오류가 발생했습니다. 잠시 후 다시 시도해주세요."

    _BUCKET = "documents-uploads"
    safe_name = name.replace(" ", "_")
    filename = f"{safe_name}_면접평가표.docx"
    file_id = str(_uuid.uuid4()).replace("-", "")
    storage_key = f"{account_id}/interview_evaluation/{file_id}/evaluation.docx"

    try:
        sb.storage.from_(_BUCKET).upload(
            path=storage_key,
            file=docx_bytes,
            file_options={
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "upsert": "false",
            },
        )
        res = sb.storage.from_(_BUCKET).create_signed_url(storage_key, expires_in=604800)
        download_url = (res or {}).get("signedURL") or (res or {}).get("signed_url")
    except Exception as _exc:
        log.exception("[evaluation_export] storage upload failed for applicant=%s: %s", name, _exc)
        return "파일 업로드 중 오류가 발생했습니다. 잠시 후 다시 시도해주세요."

    # 아티팩트 메타에 docx_url 업데이트
    try:
        updated_meta = {**meta, "docx_url": download_url}
        updated_content = artifact.get("content", "") + (
            f"\n\n---\n\n📄 **DOCX 다운로드:** [{filename}]({download_url})"
            if download_url and "DOCX 다운로드" not in artifact.get("content", "")
            else ""
        )
        sb.table("artifacts").update({
            "metadata": updated_meta,
            "content": updated_content,
        }).eq("id", artifact["id"]).execute()
    except Exception:
        pass

    return (
        f"📄 **{filename}** 파일로 저장 완료!\n\n"
        f"[다운로드 링크]({download_url})\n\n"
        "(링크는 7일간 유효합니다.)"
    )


# ──────────────────────────────────────────────────────────────────────────
# 급여 미리보기 (순수 Python 계산)
# ──────────────────────────────────────────────────────────────────────────
_PAYROLL_PREVIEW_PREFIX = "__PAYROLL_PREVIEW_REQUEST__:"
_WORK_TABLE_CONFIRMED_PREFIX = "__WORK_TABLE_CONFIRMED__:"


def _resolve_employee(account_id: str, employee_id: str | None, employee_name: str | None) -> dict:
    """UUID 또는 이름으로 직원 단건 조회. 못 찾으면 {} 반환."""
    import re as _re
    sb = get_supabase()
    _UUID_RE = _re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$", _re.I)
    if employee_id and _UUID_RE.match(employee_id):
        row = sb.table("employees").select("*").eq("id", employee_id).eq("account_id", account_id).maybe_single().execute()
        return row.data or {}
    name_query = (employee_name or employee_id or "").strip()
    if not name_query:
        return {}
    all_rows = sb.table("employees").select("*").eq("account_id", account_id).execute()
    all_emps = all_rows.data or []
    candidates = [e for e in all_emps if name_query.lower() in (e.get("name") or "").lower()]
    if len(candidates) == 1:
        return candidates[0]
    if len(candidates) > 1:
        names = ", ".join(c.get("name", "") for c in candidates)
        raise ValueError(f"'{name_query}'와 일치하는 직원이 여러 명입니다: {names}\n정확한 이름으로 다시 말씀해 주세요.")
    return {}


def _fetch_work_records(account_id: str, employee_uuid: str, month: str) -> list[dict]:
    import calendar as _cal
    sb = get_supabase()
    y, m = int(month[:4]), int(month[5:7])
    last_day = _cal.monthrange(y, m)[1]
    res = (
        sb.table("work_records")
        .select("work_date,hours_worked,overtime_hours,night_hours,holiday_hours,memo")
        .eq("employee_id", employee_uuid)
        .eq("account_id", account_id)
        .gte("work_date", f"{month}-01")
        .lte("work_date", f"{month}-{last_day:02d}")
        .order("work_date")
        .execute()
    )
    return res.data or []


def _build_payroll_reply(emp: dict, month: str, records: list[dict]) -> str:
    import json as _json
    from app.agents._payroll_calculator import calculate_payroll, format_preview_table

    emp_name = emp.get("name", "직원")
    emp_type = emp.get("employment_type", "시급제")
    hourly = int(emp.get("hourly_rate") or 0)
    monthly_sal = int(emp.get("monthly_salary") or 0)

    if not hourly and not monthly_sal:
        return f"{emp_name}의 시급 또는 월급 정보가 없어요. 직원 관리 탭에서 먼저 입력해 주세요."

    total_hours = sum(float(r.get("hours_worked", 0)) for r in records)
    total_ot = sum(float(r.get("overtime_hours", 0)) for r in records)
    total_night = sum(float(r.get("night_hours", 0)) for r in records)
    total_holiday = sum(float(r.get("holiday_hours", 0)) for r in records)

    result = calculate_payroll(
        employment_type=emp_type,
        hourly_rate=hourly,
        monthly_salary=monthly_sal,
        hours_worked=total_hours,
        overtime_hours=total_ot,
        night_hours=total_night,
        holiday_hours=total_holiday,
    )
    preview_md = format_preview_table(result, emp_name, month)

    preview_data = {
        "employee_id": emp["id"],
        "employee_name": emp_name,
        "employment_type": emp_type,
        "hourly_rate": hourly,
        "monthly_salary": monthly_sal,
        "pay_month": month,
        "pay_day": emp.get("pay_day") or 25,
        "hours_worked": total_hours,
        "overtime_hours": total_ot,
        "night_hours": total_night,
        "holiday_hours": total_holiday,
        "base_pay": result.base_pay,
        "overtime_pay": result.overtime_pay,
        "night_pay": result.night_pay,
        "holiday_pay": result.holiday_pay,
        "meal_allowance": result.meal_allowance,
        "gross_pay": result.gross_pay,
        "national_pension": result.national_pension,
        "health_insurance": result.health_insurance,
        "ltc_insurance": result.ltc_insurance,
        "employment_insurance": result.employment_insurance,
        "income_tax": result.income_tax,
        "local_income_tax": result.local_income_tax,
        "total_deductions": result.total_deductions,
        "net_pay": result.net_pay,
        "has_insurance": result.has_insurance,
    }
    return (
        f"{preview_md}\n\n"
        f"[PAYROLL_PREVIEW_DATA:{_json.dumps(preview_data, ensure_ascii=False)}]\n\n"
        "[CHOICES]\n급여명세서 생성\n취소\n[/CHOICES]"
    )


@traceable(name="recruitment.run_payroll_preview", run_type="chain")
async def run_payroll_preview(
    message: str,
    account_id: str,
    history: list | None = None,
    rag_context: str = "",
    long_term_context: str = "",
    *,
    employee_id: str | None = None,
    employee_name: str | None = None,
    pay_month: str | None = None,
    **_,
) -> str:
    import json as _json

    # ── 1. 파라미터 추출 (마커 우선) ──────────────────────────────────────────
    emp_id = employee_id
    emp_name_hint = employee_name
    month = pay_month

    # __WORK_TABLE_CONFIRMED__ 마커: Save 버튼 → 바로 급여 계산
    if _WORK_TABLE_CONFIRMED_PREFIX in message:
        try:
            raw = message[message.index(_WORK_TABLE_CONFIRMED_PREFIX) + len(_WORK_TABLE_CONFIRMED_PREFIX):]
            confirmed = _json.loads(raw.strip())
            emp_id = confirmed.get("employee_id", emp_id)
            month = confirmed.get("pay_month", month)
        except Exception:
            pass

    if _PAYROLL_PREVIEW_PREFIX in message:
        try:
            raw = message[message.index(_PAYROLL_PREVIEW_PREFIX) + len(_PAYROLL_PREVIEW_PREFIX):]
            data = _json.loads(raw.strip())
            emp_id = data.get("employee_id", emp_id)
            emp_name_hint = data.get("employee_name", emp_name_hint)
            month = data.get("pay_month", month)
        except Exception:
            pass

    # ── Step A: 직원 정보가 없으면 먼저 직원 선택 질문 ──────────────────────────
    if not emp_id and not emp_name_hint:
        try:
            _sb = get_supabase()
            _rows = (
                _sb.table("employees")
                .select("name")
                .eq("account_id", account_id)
                .order("name")
                .execute()
                .data or []
            )
            _names = [r["name"] for r in _rows if r.get("name")]
        except Exception:
            _names = []
        if _names:
            _choices = "\n".join(_names[:8])
            return (
                "누구의 급여명세서인가요?\n\n"
                f"[CHOICES]\n{_choices}\n직접 입력\n[/CHOICES]"
            )
        return "누구의 급여명세서를 만들까요? 직원 이름을 알려주세요."

    # ── 2. 직원 조회 ──────────────────────────────────────────────────────────
    try:
        emp = _resolve_employee(account_id, emp_id, emp_name_hint)
    except ValueError as e:
        return str(e)
    except Exception as _e:
        log.exception("[payroll_preview] employee lookup failed: %s", _e)
        emp = {}

    if not emp:
        return f"'{emp_name_hint or emp_id}' 직원을 찾을 수 없어요. 이름을 다시 확인해 주세요."

    employee_uuid = emp["id"]
    emp_display = emp.get("name", "직원")

    # ── Step B: 월 정보가 없으면 월 선택 질문 ────────────────────────────────
    if not month:
        from datetime import date as _date
        _today = _date.today()
        _y, _m = _today.year, _today.month
        _month_opts = []
        for _ in range(3):
            _month_opts.append(f"{_y}-{_m:02d}")
            _m -= 1
            if _m == 0:
                _m, _y = 12, _y - 1
        _choices = "\n".join(_month_opts)
        return (
            f"{emp_display}의 몇 월 급여명세서를 만들까요?\n\n"
            f"[CHOICES]\n{_choices}\n직접 입력\n[/CHOICES]"
        )

    # ── 3. Save 확인 마커 → 즉시 급여 계산 ───────────────────────────────────
    if _WORK_TABLE_CONFIRMED_PREFIX in message:
        records = []
        try:
            raw = message[message.index(_WORK_TABLE_CONFIRMED_PREFIX) + len(_WORK_TABLE_CONFIRMED_PREFIX):]
            confirmed_data = _json.loads(raw.strip())
            records = confirmed_data.get("records") or []
        except Exception:
            pass
        if not records:
            try:
                records = _fetch_work_records(account_id, employee_uuid, month)
            except Exception:
                records = []
        return _build_payroll_reply(emp, month, records)

    # ── 4. 근무 기록 조회 → 표 확인 단계 ─────────────────────────────────────
    try:
        records = _fetch_work_records(account_id, employee_uuid, month)
    except Exception:
        records = []

    emp_name = emp.get("name", "직원")

    if not records:
        _YES = {"입력하기", "ㅇㅇ", "네", "응", "예", "yes", "입력", "직접입력", "직접 입력"}
        _msg_clean = message.strip().lower().replace(" ", "")
        if any(kw.replace(" ", "") in _msg_clean for kw in _YES):
            empty_payload = _json.dumps(
                {"employee_id": employee_uuid, "employee_name": emp_name, "pay_month": month, "records": []},
                ensure_ascii=False,
            )
            return (
                f"{emp_name} {month} 근무 기록 입력 화면을 열어드릴게요.\n\n"
                f"[ACTION:OPEN_WORK_TABLE:{empty_payload}]"
            )
        return (
            f"{emp_name}의 {month} 근무 기록이 없어요. 직접 입력하시겠어요?\n\n"
            "[CHOICES]\n입력하기\n취소\n[/CHOICES]"
        )

    # 근무 기록을 표 액션 마커로 반환
    work_table_payload = {
        "employee_id": employee_uuid,
        "employee_name": emp_name,
        "pay_month": month,
        "records": [
            {
                "work_date": r["work_date"],
                "hours_worked": float(r.get("hours_worked", 0)),
                "overtime_hours": float(r.get("overtime_hours", 0)),
                "night_hours": float(r.get("night_hours", 0)),
                "holiday_hours": float(r.get("holiday_hours", 0)),
                "memo": r.get("memo") or "",
            }
            for r in records
        ],
    }
    return (
        f"{emp_name}의 {month} 근무 기록입니다. 확인 후 저장하면 급여명세서를 계산해 드릴게요.\n\n"
        f"[ACTION:OPEN_WORK_TABLE:{_json.dumps(work_table_payload, ensure_ascii=False)}]"
    )


def describe(account_id: str) -> list[dict]:
    """OpenAI tools 스펙용 capability 매니페스트."""
    caps: list[dict] = [
        {
            "name": "recruit_posting_set",
            "description": (
                "당근알바·알바천국·사람인 3종 플랫폼 채용공고를 하나의 카드로 작성한다. "
                "모든 필수 정보(매장명·급여·근무지·고용형태·모집인원·근무요일·근무시각·모집기간)가 "
                "확정된 경우에만 호출. 하나라도 미확정이면 [CHOICES] 로 되묻는다."
            ),
            "handler": run_posting_set,
            "parameters": {
                "type": "object",
                "properties": {
                    "position":        {"type": "string", "description": "직종 또는 포지션명 (예: 바리스타, 홀서빙)"},
                    "business_name":   {"type": "string", "description": "매장명/상호/가게 이름"},
                    "location":        {"type": "string", "description": "근무지/매장 주소"},
                    "employment_type": {"type": "string", "enum": ["정규직", "계약직", "파트타임", "알바", "단기"]},
                    "headcount":       {"type": "integer", "description": "모집 인원"},
                    "work_days":       {"type": "array", "items": {"type": "string"}, "description": "근무 요일 (예: ['월','화','수'])"},
                    "work_start":      {"type": "string", "description": "근무 시작 시각 HH:MM"},
                    "work_end":        {"type": "string", "description": "근무 종료 시각 HH:MM"},
                    "wage_hourly":     {"type": "integer", "description": "시급(원). 최저임금 이상"},
                    "wage_monthly":    {"type": "integer", "description": "월급(원)"},
                    "annual_salary":   {"type": "integer", "description": "연봉(원)"},
                    "start_date":      {"type": "string", "description": "모집 시작일 YYYY-MM-DD"},
                    "end_date":        {"type": "string", "description": "모집 마감일 YYYY-MM-DD"},
                    "weekly_hours":    {"type": "number", "description": "주 근무시간 (선택)"},
                    "extra_note":      {"type": "string", "description": "자유 기술 (선택)"},
                },
                "required": [
                    "position", "business_name", "location", "employment_type",
                    "headcount", "work_days", "work_start", "work_end", "start_date", "end_date",
                ],
            },
        },
        {
            "name": "recruit_interview",
            "description": "특정 직종·레벨 지원자용 면접 질문 세트를 생성한다.",
            "handler": run_interview,
            "parameters": {
                "type": "object",
                "properties": {
                    "position": {"type": "string"},
                    "level":    {"type": "string", "enum": ["신입", "경력", "시니어"], "default": "신입"},
                    "count":    {"type": "integer", "default": 5, "minimum": 3, "maximum": 20},
                },
                "required": ["position"],
            },
        },
        {
            "name": "recruit_hiring_drive",
            "description": (
                "공채·시즌 채용 기간을 artifact 로 등록한다 (스케쥴러가 D-7/3/1/0 리마인드 자동 발사). "
                "반드시 시작일·종료일·모집 인원이 확정된 경우에만 호출."
            ),
            "handler": run_hiring_drive,
            "parameters": {
                "type": "object",
                "properties": {
                    "title":           {"type": "string", "description": "채용 기간 이름 (예: '2026년 상반기 공채')"},
                    "start_date":      {"type": "string", "description": "시작일 YYYY-MM-DD"},
                    "end_date":        {"type": "string", "description": "종료일 YYYY-MM-DD"},
                    "headcount":       {"type": "integer"},
                    "target_position": {"type": "string", "description": "대상 직종(선택)"},
                },
                "required": ["title", "start_date", "end_date", "headcount"],
            },
        },
        {
            "name": "recruit_checklist_guide",
            "description": "채용 관련 체크리스트 또는 가이드 문서를 작성한다.",
            "handler": run_checklist_guide,
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "예: '근로계약서 체크리스트'"},
                    "kind":  {"type": "string", "enum": ["checklist", "guide"], "default": "checklist"},
                },
                "required": ["topic"],
            },
        },
        {
            "name": "recruit_onboarding",
            "description": (
                "신규 입사자 온보딩 자료를 생성한다. "
                "체크리스트(입사 단계별 할 일) / 플랜(수습 기간 일정) / 교육자료(매뉴얼·SOP) 중 "
                "사용자 요청에 맞는 타입으로 저장. Onboarding 서브허브에 자동 분류."
            ),
            "handler": run_onboarding,
            "parameters": {
                "type": "object",
                "properties": {
                    "position": {
                        "type": "string",
                        "description": "직책 또는 직종 (예: 바리스타, 홀서빙, 주방장)",
                    },
                    "onboarding_type": {
                        "type": "string",
                        "enum": ["onboarding_checklist", "onboarding_plan", "education_material"],
                        "description": "자료 타입. 불명확하면 생략 — 에이전트가 감지하거나 [CHOICES] 로 확인.",
                    },
                },
                "required": ["position"],
            },
        },
        {
            "name": "recruit_payroll_preview",
            "description": (
                "직원 DB에서 근무 기록을 읽어 급여명세서 미리보기를 계산한다. "
                "순수 Python 수식으로 4대보험·소득세 공제액을 산출하고 표 형식으로 보여준다. "
                "사용자가 확인하면 Documents 에이전트가 Excel을 생성한다. "
                "급여/월급/페이/봉급/명세서 키워드가 있을 때 호출."
            ),
            "handler": run_payroll_preview,
            "parameters": {
                "type": "object",
                "properties": {
                    "employee_id": {"type": "string", "description": "직원 UUID. UUID를 모르면 생략하고 employee_name을 사용."},
                    "employee_name": {"type": "string", "description": "직원 이름 (예: 송진우). employee_id UUID를 모를 때 사용."},
                    "pay_month": {"type": "string", "description": "급여 정산 월 YYYY-MM (예: 2026-04)"},
                },
                "required": [],
            },
        },
        {
            "name": "recruit_resume_parse",
            "description": (
                "구직자 이력서 파일을 파싱해 DB에 저장한다. "
                "사용자가 이력서 파일을 업로드하고 파싱/분석을 요청할 때 호출. "
                "upload_payload 또는 upload_payloads contextvar 에 파일 내용이 있어야 한다. "
                "⚠️ 사용자 메시지에 면접/질문 키워드가 있으면 파싱 후 내부에서 면접 질문까지 직접 생성하므로 "
                "recruit_resume_interview 를 별도 dispatch 하지 말 것."
            ),
            "handler": run_resume_parse,
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
        {
            "name": "recruit_resume_interview",
            "description": (
                "저장된 이력서를 바탕으로 날카로운 맞춤 면접 질문을 생성하고 artifact 로 저장한다. "
                "이력서 파싱 완료 후 특정 지원자의 면접 질문을 요청할 때 호출."
            ),
            "handler": run_resume_interview,
            "parameters": {
                "type": "object",
                "properties": {
                    "applicant_name": {
                        "type": "string",
                        "description": "면접 질문을 생성할 지원자 이름 (이력서에서 파싱된 이름)",
                    },
                    "count": {
                        "type": "integer",
                        "description": "생성할 면접 질문 수 (기본 7)",
                        "default": 7,
                        "minimum": 3,
                        "maximum": 15,
                    },
                },
                "required": ["applicant_name"],
            },
        },
        {
            "name": "recruit_interview_evaluation",
            "description": (
                "저장된 이력서를 바탕으로 지원자 맞춤 면접 평가표를 생성하고 캔버스에 마크다운으로 표시한다. "
                "상단 종합 점수표(배점·채점란)와 하단 역량별 평가 기준·코멘트란으로 구성. "
                "평가 항목·배점 비율 커스터마이징 지원. Interviews 서브허브에 자동 분류. "
                "사용자가 '면접 평가표/평가 시트/평가 양식' 을 요청할 때 호출. "
                "⚠️ DOCX 파일 저장은 recruit_evaluation_export_docx 를 별도 호출."
            ),
            "handler": run_interview_evaluation,
            "parameters": {
                "type": "object",
                "properties": {
                    "applicant_name": {
                        "type": "string",
                        "description": "평가표를 만들 지원자 이름 (이력서에서 파싱된 이름)",
                    },
                    "position": {
                        "type": "string",
                        "description": "지원 직종 (선택 — 이력서에 희망직종이 있으면 생략 가능)",
                    },
                    "custom_categories": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "커스텀 평가 항목 목록 (생략 시 기본 3항목 사용)",
                    },
                    "weights": {
                        "type": "object",
                        "description": "항목별 배점 비율 (예: {\"기술역량\": 50, \"태도\": 30, \"소통\": 20}, 합계 100)",
                        "additionalProperties": {"type": "integer"},
                    },
                },
                "required": ["applicant_name"],
            },
        },
        {
            "name": "recruit_evaluation_export_docx",
            "description": (
                "캔버스에 표시된 면접 평가표를 DOCX 파일로 변환해 다운로드 URL을 반환한다. "
                "사용자가 평가표 검토·수정 후 'DOCX로 저장/내보내기/다운로드' 를 요청할 때 호출. "
                "artifact_id 미확정 시 최근 interview_evaluation artifact 를 자동 사용."
            ),
            "handler": run_evaluation_export_docx,
            "parameters": {
                "type": "object",
                "properties": {
                    "artifact_id": {
                        "type": "string",
                        "description": "DOCX로 변환할 interview_evaluation artifact ID (선택 — 미확정 시 생략)",
                    },
                    "applicant_name": {
                        "type": "string",
                        "description": "지원자 이름으로 artifact 검색 시 사용 (선택)",
                    },
                },
                "required": [],
            },
        },
    ]

    # 포스터 capability 는 posting_set 이 있을 때만 노출
    if _find_recent_posting_set(account_id):
        caps.append({
            "name": "recruit_posting_poster",
            "description": (
                "저장된 채용공고 세트를 선택해 GPT-4o 로 standalone HTML 포스터를 생성한다. "
                "플랫폼 복수 선택 가능(당근알바·알바천국·사람인). "
                "posting_set_id 미확정 시 목록을 보여주고, platforms 미확정 시 선택을 요청한다. "
                "사용자가 '이미지/포스터/배너/썸네일' 을 요청할 때만 호출."
            ),
            "handler": run_posting_poster,
            "parameters": {
                "type": "object",
                "properties": {
                    "posting_set_id": {"type": "string", "description": "포스터를 생성할 job_posting_set artifact ID. 미확정이면 생략."},
                    "platforms":      {
                        "type": "array",
                        "items": {"type": "string", "enum": list(VALID_PLATFORMS)},
                        "description": "포스터를 만들 플랫폼 목록 (복수 가능). 미확정이면 생략.",
                    },
                    "style":          {"type": "string", "description": "자유 디자인 지시 (예: '따뜻한 브라운 톤, 미니멀')"},
                },
            },
        })

    return caps


# ──────────────────────────────────────────────────────────────────────────
# 메인 run (legacy fallback 겸 capability wrapper 타겟)
# ──────────────────────────────────────────────────────────────────────────
@traceable(name="recruitment.run", run_type="chain")
async def run(
    message: str,
    account_id: str,
    history: list[dict],
    rag_context: str = "",
    long_term_context: str = "",
) -> str:
    # 1. 의도 휴리스틱
    type_guess, cat_hint = detect_recruit_intent(message)
    if not type_guess:
        for h in reversed(history[-6:]):
            if h.get("role") == "user":
                t2, c2 = detect_recruit_intent(h.get("content") or "")
                if t2:
                    type_guess = t2
                    cat_hint = cat_hint or c2
                    break

    # 업종: profile 우선 → 메시지 힌트 fallback
    business_type = _get_business_type(account_id)
    if not business_type and cat_hint:
        business_type = cat_hint

    want_posting = type_guess in ("job_posting", "job_posting_set")
    want_image = wants_posting_poster(message)

    recruit_ctx = build_recruit_context(
        business_type=business_type,
        want_job_posting=want_posting,
        want_image=want_image,
    )

    system = SYSTEM_PROMPT_BASE + "\n\n" + today_context() + "\n\n" + recruit_ctx

    hubs = list_sub_hub_titles(account_id, "recruitment")
    if hubs:
        system += "\n\n[이 계정의 recruitment 서브허브]\n- " + "\n- ".join(hubs)
    if long_term_context:
        system += f"\n\n[사용자 장기 기억]\n{long_term_context}"
    if rag_context:
        system += f"\n\n{rag_context}"
    fb = feedback_context(account_id, "recruitment")
    if fb:
        system += f"\n\n{fb}"

    resp = await chat_completion(
        messages=[
            {"role": "system", "content": system},
            *history,
            {"role": "user", "content": message},
        ],
    )
    reply = resp.choices[0].message.content or ""

    # 2. 3종 공고 마커 → 세트 + 3개 자식 저장
    reply = await _maybe_dispatch_posting_set(account_id, reply)

    # 3. 포스터 마커 → GPT-4o 로 HTML 포스터 생성 + Storage(html) + artifact 저장
    reply = await _maybe_dispatch_poster(account_id, reply)

    # 4. 일반 [ARTIFACT] (interview_questions / hiring_drive / checklist / guide / 단일 job_posting)
    await save_artifact_from_reply(
        account_id,
        "recruitment",
        reply,
        default_title="채용 자료",
        valid_types=_GENERIC_VALID_TYPES,
        extra_meta_keys=("due_label",),
    )
    return reply
